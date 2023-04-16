import sys
import re
import os
import json
import docx

from logging import warning
from argparse import ArgumentParser


def argparser():
    ap = ArgumentParser()
    ap.add_argument('jsonl')
    ap.add_argument('outdir')
    ap.add_argument('--max-len', type=int, default=950000)
    return ap


def yield_data(fn):
    with open(fn) as f:
        for l in f:
            yield json.loads(l)


def clean(text):
    # Remove NULL and control chars, keep newlines and tabs
    text = re.sub(r'[\x00-\x08\x0b-\x1f\x7f-\x9f]', '', text)
    return text


def main(argv):
    args = argparser().parse_args(argv[1:])

    root = os.path.splitext(os.path.basename(args.jsonl))[0]
    
    doc = docx.Document()
    doc_count = 0
    current_len = 0

    for d_idx, d in enumerate(yield_data(args.jsonl)):
        text = d['text']
        
        cleaned = clean(text)
        if cleaned != text:
            warning(f'clean() modified {d["message_id"]}')
            text = cleaned
        
        d_len = len(text) + 50
        if current_len + d_len > args.max_len:
            doc.save(f'{args.outdir}/{root}-{doc_count:05d}.docx')
            doc_count += 1
            current_len = 0
            doc = docx.Document()

        pgraph = doc.add_paragraph('')

        r = pgraph.add_run(f'Document {d_idx}')
        r.bold = True
        r.underline = True

        pgraph = doc.add_paragraph(text)

        current_len += d_len
    else:
        doc.save(f'{args.outdir}/{root}-{doc_count:05d}.docx')


if __name__=='__main__':
    sys.exit(main(sys.argv))
